#!/usr/bin/env python3
"""
PDF → DOCX converter using PyMuPDF + python-docx.

No OpenCV / libGL dependency — replaces the previous pdf2docx approach which
failed in Lambda because pdf2docx imports cv2 at import time and cv2 links
against libGL.so.1 (OpenGL), which is absent in the Lambda container image.

Strategy:
  - Text-based pages (≥30 extractable chars): extract blocks via PyMuPDF,
    write paragraphs to DOCX preserving rough block structure.
  - Scanned/image-only pages: OCR via Tesseract (if installed), then write
    OCR'd lines as paragraphs.  If Tesseract is unavailable or OCR fails,
    embed the rendered page as an inline PNG image so output is never blank.

Usage: python3 pdf_to_docx.py <input.pdf> <output.docx>
Exit 0 = success (prints "OK:<bytes>" to stdout).
Exit 1 = failure (error message on stderr).
"""

import sys
import os
import subprocess
import tempfile
import shutil

MIN_TEXT_CHARS = 30   # chars per page; below this we treat the page as scanned
MIN_DOCX_BYTES = 3000  # a valid non-empty DOCX is always larger than this


# ---------------------------------------------------------------------------
# OCR helpers
# ---------------------------------------------------------------------------

def _has_tesseract() -> bool:
    return shutil.which("tesseract") is not None


def _ocr_page(page, dpi: int = 300) -> str:
    """Render *page* (fitz.Page) to a PNG at *dpi* and OCR it with Tesseract."""
    import fitz  # noqa: PLC0415
    mat = fitz.Matrix(dpi / 72, dpi / 72)
    pix = page.get_pixmap(matrix=mat, alpha=False)
    img_data = pix.tobytes("png")

    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
        tmp.write(img_data)
        img_path = tmp.name

    try:
        result = subprocess.run(
            [
                "tesseract", img_path, "stdout",
                "-l", "eng",
                "--psm", "3",   # fully automatic page segmentation
                "--oem", "3",   # default OCR engine (LSTM + legacy)
            ],
            capture_output=True,
            text=True,
            timeout=120,
        )
        if result.returncode != 0:
            raise RuntimeError(
                f"tesseract exited {result.returncode}: {result.stderr.strip()[:300]}"
            )
        return result.stdout
    finally:
        try:
            os.unlink(img_path)
        except OSError:
            pass


# ---------------------------------------------------------------------------
# Image-embedding fallback
# ---------------------------------------------------------------------------

def _embed_page_as_image(doc_out, page, dpi: int = 150) -> None:
    """Render *page* as a PNG and embed it as an inline picture in *doc_out*."""
    import fitz  # noqa: PLC0415
    from docx.shared import Inches  # noqa: PLC0415

    mat = fitz.Matrix(dpi / 72, dpi / 72)
    pix = page.get_pixmap(matrix=mat, alpha=False)
    img_data = pix.tobytes("png")
    # Fit to a 6-inch content-area width (standard letter margins)
    width_in = min(6.0, pix.width / dpi)

    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
        tmp.write(img_data)
        tmp_path = tmp.name

    try:
        doc_out.add_picture(tmp_path, width=Inches(width_in))
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


# ---------------------------------------------------------------------------
# Core conversion
# ---------------------------------------------------------------------------

def convert(input_path: str, output_path: str) -> None:
    import fitz  # noqa: PLC0415
    from docx import Document  # noqa: PLC0415
    from docx.shared import Pt  # noqa: PLC0415

    doc_in = fitz.open(input_path)
    n_pages = len(doc_in)
    if n_pages == 0:
        raise ValueError("PDF has no pages")

    use_ocr = _has_tesseract()
    doc_out = Document()

    # Tighten default paragraph spacing for a cleaner result
    doc_out.styles["Normal"].paragraph_format.space_after = Pt(4)

    for page_idx in range(n_pages):
        page = doc_in[page_idx]

        if page_idx > 0:
            doc_out.add_page_break()

        text = page.get_text("text").strip()

        if len(text) >= MIN_TEXT_CHARS:
            # Text-based page — write block-level paragraphs
            blocks = page.get_text("blocks")
            # blocks: (x0, y0, x1, y1, text, block_no, block_type)
            # block_type 0 = text, 1 = image
            for block in blocks:
                if len(block) < 7 or block[6] != 0:
                    continue
                block_text = block[4].strip()
                if not block_text:
                    continue
                para = doc_out.add_paragraph(block_text)
                para.paragraph_format.space_after = Pt(6)

        elif use_ocr:
            # Scanned page — OCR, fall back to image on any failure
            try:
                ocr_text = _ocr_page(page)
                if ocr_text.strip():
                    for line in ocr_text.splitlines():
                        stripped = line.strip()
                        if stripped:
                            doc_out.add_paragraph(stripped)
                else:
                    # Tesseract found no text → embed as image
                    _embed_page_as_image(doc_out, page)
            except Exception:  # pylint: disable=broad-except
                _embed_page_as_image(doc_out, page)

        else:
            # Tesseract not installed → embed rendered page as image
            _embed_page_as_image(doc_out, page)

    doc_in.close()
    doc_out.save(output_path)

    if not os.path.exists(output_path):
        raise RuntimeError("Output file was not created")

    size = os.path.getsize(output_path)
    if size < MIN_DOCX_BYTES:
        raise RuntimeError(
            f"Output DOCX is suspiciously small ({size} bytes); "
            "the file may be blank or corrupt"
        )


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main() -> None:
    if len(sys.argv) != 3:
        print(f"Usage: {sys.argv[0]} <input.pdf> <output.docx>", file=sys.stderr)
        sys.exit(1)

    input_path, output_path = sys.argv[1], sys.argv[2]

    if not os.path.exists(input_path):
        print(f"Input file not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    # Dependency check — give a clear error if packages are missing
    try:
        import fitz  # noqa: F401, PLC0415
    except ImportError as exc:
        print(f"PyMuPDF not available: {exc}", file=sys.stderr)
        sys.exit(1)

    try:
        from docx import Document  # noqa: F401, PLC0415
    except ImportError as exc:
        print(f"python-docx not available: {exc}", file=sys.stderr)
        sys.exit(1)

    try:
        convert(input_path, output_path)
        size = os.path.getsize(output_path)
        print(f"OK:{size}", flush=True)
        sys.exit(0)
    except Exception as exc:  # pylint: disable=broad-except
        print(f"Conversion failed: {exc}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
